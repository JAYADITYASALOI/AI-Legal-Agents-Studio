from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass(frozen=True)
class LoadedDocument:
    title: str
    path: Path
    text: str


def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore").strip()


def _read_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n".join(paragraphs).strip()


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            pages.append(text.strip())
    return "\n\n".join(pages).strip()


def extract_text_from_file(file_path: Path | str | None) -> str:
    if not file_path:
        return ""
    path = Path(file_path)
    if not path.exists():
        return ""
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return _read_text_file(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    try:
        return _read_text_file(path)
    except Exception:
        return ""


def load_documents(paths: Iterable[Path | str]) -> list[LoadedDocument]:
    docs: list[LoadedDocument] = []
    for raw in paths:
        path = Path(raw)
        if not path.exists() or path.is_dir():
            continue
        text = extract_text_from_file(path)
        if not text.strip():
            continue
        docs.append(LoadedDocument(title=path.name, path=path, text=text.strip()))
    return docs
